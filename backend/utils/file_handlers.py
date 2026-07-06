import io
import re
import json
import chardet
import aiofiles
import pandas as pd
import polars as pl
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from loguru import logger


SUPPORTED_EXTENSIONS = {
    ".csv", ".xlsx", ".xls", ".json", ".xml",
    ".pdf", ".docx", ".doc", ".txt", ".tsv", ".parquet"
}


async def detect_encoding(file_path: str) -> str:
    async with aiofiles.open(file_path, "rb") as f:
        raw = await f.read(100_000)
    result = chardet.detect(raw)
    return result.get("encoding", "utf-8") or "utf-8"


async def read_csv(file_path: str) -> pd.DataFrame:
    encoding = await detect_encoding(file_path)
    # Detect the separator on a small sample instead of parsing the full file per candidate
    sep = ","
    for candidate in [",", ";", "\t", "|"]:
        try:
            sample = pd.read_csv(file_path, encoding=encoding, sep=candidate, nrows=50)
            if len(sample.columns) > 1:
                sep = candidate
                break
        except Exception:
            continue
    return pd.read_csv(file_path, encoding=encoding, sep=sep, low_memory=False)


async def read_excel(file_path: str) -> pd.DataFrame:
    xl = pd.ExcelFile(file_path)
    if len(xl.sheet_names) == 1:
        return pd.read_excel(file_path, sheet_name=0)
    # Merge all sheets with a source column
    dfs = []
    for sheet in xl.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet)
        df["__sheet__"] = sheet
        dfs.append(df)
    return pd.concat(dfs, ignore_index=True)


async def read_json(file_path: str) -> pd.DataFrame:
    async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
        data = json.loads(await f.read())
    if isinstance(data, list):
        return pd.json_normalize(data)
    elif isinstance(data, dict):
        # Try to find a list inside
        for val in data.values():
            if isinstance(val, list) and len(val) > 0:
                return pd.json_normalize(val)
        return pd.json_normalize([data])
    return pd.DataFrame()


async def read_pdf(file_path: str) -> pd.DataFrame:
    try:
        import pdfplumber
        all_tables = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        all_tables.append(df)
        if all_tables:
            return pd.concat(all_tables, ignore_index=True)
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}")
    try:
        import fitz  # PyMuPDF
        text_data = []
        doc = fitz.open(file_path)
        for page in doc:
            text_data.append(page.get_text())
        doc.close()
        combined_text = "\n".join(text_data)
        return pd.DataFrame({"extracted_text": [combined_text]})
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return pd.DataFrame()


async def read_docx(file_path: str) -> pd.DataFrame:
    from docx import Document
    doc = Document(file_path)
    tables_data = []
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        rows = []
        for row in table.rows[1:]:
            rows.append([cell.text.strip() for cell in row.cells])
        if headers and rows:
            df = pd.DataFrame(rows, columns=headers)
            tables_data.append(df)
    if tables_data:
        return pd.concat(tables_data, ignore_index=True)
    # Fall back to paragraph extraction
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return pd.DataFrame({"content": paragraphs})


async def read_parquet(file_path: str) -> pd.DataFrame:
    return pd.read_parquet(file_path)


async def read_xml(file_path: str) -> pd.DataFrame:
    return pd.read_xml(file_path)


async def load_file(file_path: str) -> Tuple[pd.DataFrame, str]:
    path = Path(file_path)
    ext = path.suffix.lower()
    logger.info(f"Loading file: {file_path} (ext: {ext})")

    handlers = {
        ".csv": read_csv,
        ".tsv": read_csv,
        ".txt": read_csv,
        ".xlsx": read_excel,
        ".xls": read_excel,
        ".json": read_json,
        ".pdf": read_pdf,
        ".docx": read_docx,
        ".doc": read_docx,
        ".parquet": read_parquet,
        ".xml": read_xml,
    }

    handler = handlers.get(ext)
    if not handler:
        raise ValueError(f"Unsupported file type: {ext}")

    df = await handler(file_path)
    return df, ext


def clean_column_names(df: pd.DataFrame) -> pd.DataFrame:
    df.columns = [
        re.sub(r"[^a-zA-Z0-9_]", "_", str(col).strip())
        .strip("_")
        .replace("__", "_")
        for col in df.columns
    ]
    return df


def infer_and_cast_types(df: pd.DataFrame) -> pd.DataFrame:
    for col in df.select_dtypes(include=[object]).columns:
        non_null = df[col].dropna()
        if non_null.empty:
            continue
        # Probe a sample first — only convert the full column if the sample passes
        sample = non_null.sample(n=min(len(non_null), 500), random_state=0)

        numeric_sample = pd.to_numeric(sample, errors="coerce")
        if numeric_sample.notna().mean() > 0.8:
            numeric_series = pd.to_numeric(df[col], errors="coerce")
            if numeric_series.notna().sum() / max(len(df), 1) > 0.8:
                df[col] = numeric_series
            continue

        try:
            date_sample = pd.to_datetime(sample, errors="coerce", format="mixed")
            if date_sample.notna().mean() > 0.8:
                date_series = pd.to_datetime(df[col], errors="coerce", format="mixed")
                if date_series.notna().sum() / max(len(df), 1) > 0.8:
                    df[col] = date_series
        except Exception:
            pass
    return df


def dataframe_to_dict(df: pd.DataFrame, max_rows: int = 10000) -> Dict[str, Any]:
    if len(df) > max_rows:
        df = df.head(max_rows)
    return {
        "columns": list(df.columns),
        "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
        "shape": list(df.shape),
        "data": df.to_dict(orient="records"),
    }


async def load_from_url(url: str) -> pd.DataFrame:
    import asyncio
    import requests
    # requests is blocking — run it in a thread so the event loop stays responsive
    response = await asyncio.to_thread(requests.get, url, timeout=30)
    response.raise_for_status()

    content_type = response.headers.get("content-type", "")
    if "csv" in content_type or url.endswith(".csv"):
        return pd.read_csv(io.StringIO(response.text))
    elif "json" in content_type or url.endswith(".json"):
        data = response.json()
        if isinstance(data, list):
            return pd.json_normalize(data)
        return pd.json_normalize([data])
    elif "excel" in content_type or url.endswith((".xlsx", ".xls")):
        return pd.read_excel(io.BytesIO(response.content))
    else:
        return pd.read_csv(io.StringIO(response.text))
